"""Bring existing notes in from other file types.

A student's material rarely starts as audio: there are lecture slides as PDFs,
a handout in Word, a photo of a whiteboard. This module turns those into plain
text so they can run through exactly the same template → summary pipeline as a
recording, and be filed alongside everything else.

Every dependency here is optional and imported lazily. A missing library (or a
missing Tesseract binary for image OCR) degrades to a clear, actionable message
rather than breaking the app or the other formats.
"""
import codecs
import io
import re
from pathlib import Path

TEXT_SUFFIXES = {".txt", ".md", ".markdown", ".csv"}
RTF_SUFFIXES = {".rtf"}
PDF_SUFFIXES = {".pdf"}
DOCX_SUFFIXES = {".docx"}
IMAGE_SUFFIXES = {".png", ".jpg", ".jpeg", ".webp", ".tif", ".tiff", ".bmp"}

SUPPORTED_SUFFIXES = (TEXT_SUFFIXES | RTF_SUFFIXES | PDF_SUFFIXES | DOCX_SUFFIXES
                      | IMAGE_SUFFIXES)

# For st.file_uploader, which wants extensions without the dot.
UPLOAD_TYPES = sorted(suffix.lstrip(".") for suffix in SUPPORTED_SUFFIXES)

# Enough characters that we believe we got a real text layer rather than a
# handful of stray glyphs from a scanned page.
_MIN_MEANINGFUL_CHARS = 20


class ImportFailed(Exception):
    """Raised when a file can't be read at all, with a message meant for the user."""


def _decode_text(data: bytes) -> str:
    """Decode a text file, guessing the encoding conservatively.

    UTF-16 is only attempted when a byte-order mark says so: it happily decodes
    almost any even-length byte string, so trying it speculatively turns
    perfectly good cp1252 text into mojibake instead of failing over.
    """
    if data.startswith((codecs.BOM_UTF16_LE, codecs.BOM_UTF16_BE)):
        try:
            return data.decode("utf-16")
        except UnicodeDecodeError:
            pass
    if data.startswith(codecs.BOM_UTF8):
        return data.decode("utf-8-sig", errors="replace")
    for encoding in ("utf-8", "cp1252", "latin-1"):  # latin-1 never raises
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


_RTF_IGNORABLE = re.compile(r"\{\\\*(?:[^{}]|\{[^{}]*\})*\}")
_RTF_UNICODE = re.compile(r"\\u(-?\d+)\s?\??")
_RTF_HEX = re.compile(r"\\'([0-9a-fA-F]{2})")
_RTF_CONTROL = re.compile(r"\\([a-zA-Z]+)(-?\d+)? ?")
_RTF_TABLES = ("fonttbl", "colortbl", "stylesheet", "listtable", "listoverridetable",
               "info", "pntext", "generator", "themedata", "datastore")


def _extract_rtf(data: bytes):
    """Pull the readable text out of an RTF file.

    Without this, an `.rtf` import would hand the model a page of `\\rtf1\\ansi`
    control words as if they were the student's notes.
    """
    text = _decode_text(data)
    if not text.lstrip().startswith("{\\rtf"):
        return text.strip(), []  # mislabelled: it's really plain text

    for table in _RTF_TABLES:            # font/colour/style tables carry no prose
        text = re.sub(r"\{\\" + table + r"(?:[^{}]|\{[^{}]*\})*\}", "", text)
    text = _RTF_IGNORABLE.sub("", text)  # \* destinations are skippable by spec

    text = re.sub(r"\\(?:par|line|sect|page)\b ?", "\n", text)
    text = re.sub(r"\\tab\b ?", "\t", text)

    text = _RTF_UNICODE.sub(
        lambda m: chr(int(m.group(1)) % 65536) if m.group(1) else "", text)
    text = _RTF_HEX.sub(
        lambda m: bytes([int(m.group(1), 16)]).decode("cp1252", "replace"), text)

    # Shield the escaped literals before stripping the remaining control words
    text = (text.replace("\\\\", "\x00").replace("\\{", "\x01").replace("\\}", "\x02"))
    text = _RTF_CONTROL.sub("", text)
    text = text.replace("{", "").replace("}", "")
    text = (text.replace("\x00", "\\").replace("\x01", "{").replace("\x02", "}"))

    lines = [line.strip() for line in text.split("\n")]
    cleaned = "\n".join(line for line in lines if line)
    if not cleaned.strip():
        return "", ["No readable text could be extracted from that RTF file."]
    return cleaned.strip(), []


def _extract_pdf(data: bytes):
    try:
        from pypdf import PdfReader
    except ImportError:
        raise ImportFailed(
            "Reading PDFs needs the `pypdf` package. Install it with "
            "`pip install pypdf` (or use the Docker launcher, which includes it)."
        )

    warnings = []
    try:
        reader = PdfReader(io.BytesIO(data))
    except Exception as exc:  # pypdf raises several unrelated types on bad input
        raise ImportFailed(f"That PDF couldn't be opened ({exc}).")

    if getattr(reader, "is_encrypted", False):
        try:
            reader.decrypt("")  # many "encrypted" PDFs have an empty owner password
        except Exception:
            raise ImportFailed(
                "That PDF is password-protected. Remove the password and try again."
            )

    pages = []
    for number, page in enumerate(reader.pages, start=1):
        try:
            pages.append(page.extract_text() or "")
        except Exception:
            warnings.append(f"Page {number} couldn't be read and was skipped.")
    text = "\n\n".join(p.strip() for p in pages if p.strip())

    if len(text) < _MIN_MEANINGFUL_CHARS:
        warnings.append(
            f"This PDF has {len(reader.pages)} page(s) but almost no selectable text, "
            "so it's probably a scan. Export the pages as images and import those "
            "instead — EchoPad will run OCR on them."
        )
    return text, warnings


def _extract_docx(data: bytes):
    try:
        import docx  # python-docx
    except ImportError:
        raise ImportFailed(
            "Reading Word documents needs the `python-docx` package. Install it "
            "with `pip install python-docx` (or use the Docker launcher)."
        )
    try:
        document = docx.Document(io.BytesIO(data))
    except Exception as exc:
        raise ImportFailed(
            f"That Word file couldn't be opened ({exc}). Note that older `.doc` "
            "files aren't supported — re-save it as `.docx` first."
        )

    blocks = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                blocks.append(" | ".join(cells))
    return "\n\n".join(blocks), []


def _extract_image(data: bytes):
    try:
        import pytesseract
        from PIL import Image
    except ImportError:
        raise ImportFailed(
            "Reading text out of images needs `pytesseract` and `Pillow`. Install "
            "them with `pip install pytesseract Pillow`, plus the Tesseract OCR "
            "engine itself (macOS: `brew install tesseract`, Ubuntu: "
            "`sudo apt install tesseract-ocr`, Windows: see the Tesseract docs)."
        )
    try:
        image = Image.open(io.BytesIO(data))
    except Exception as exc:
        raise ImportFailed(f"That image couldn't be opened ({exc}).")

    try:
        text = pytesseract.image_to_string(image)
    except Exception:
        raise ImportFailed(
            "The Tesseract OCR engine isn't installed or isn't on PATH. Install it "
            "(macOS: `brew install tesseract`, Ubuntu: `sudo apt install "
            "tesseract-ocr`) and restart EchoPad. The Docker launcher includes it."
        )

    text = text.strip()
    warnings = []
    if len(text) < _MIN_MEANINGFUL_CHARS:
        warnings.append(
            "Very little text came out of this image. OCR handles clear printed "
            "text well but struggles with handwriting, glare, and angled photos — "
            "a straight-on, well-lit photo usually helps."
        )
    else:
        warnings.append(
            "Text came from OCR, so check it for misread characters before relying on it."
        )
    return text, warnings


def extract_text(filename: str, data: bytes):
    """Extract text from an uploaded file.

    Returns (text, warnings). Raises ImportFailed with a user-facing message if
    the format is unsupported or the file can't be read.
    """
    suffix = Path(filename or "").suffix.lower()

    if suffix in TEXT_SUFFIXES:
        return _decode_text(data).strip(), []
    if suffix in RTF_SUFFIXES:
        return _extract_rtf(data)
    if suffix in PDF_SUFFIXES:
        return _extract_pdf(data)
    if suffix in DOCX_SUFFIXES:
        return _extract_docx(data)
    if suffix in IMAGE_SUFFIXES:
        return _extract_image(data)

    if suffix == ".doc":
        raise ImportFailed("Older `.doc` files aren't supported — re-save as `.docx` first.")
    if suffix in (".pptx", ".ppt", ".key"):
        raise ImportFailed(
            "Slide decks aren't supported directly. Export the slides to PDF and import that."
        )
    raise ImportFailed(
        f"EchoPad doesn't know how to read `{suffix or filename}`. Supported: "
        + ", ".join(sorted(SUPPORTED_SUFFIXES))
    )


def describe_source(filename: str) -> str:
    """Short human label for where an imported note's text came from."""
    suffix = Path(filename or "").suffix.lower()
    if suffix in IMAGE_SUFFIXES:
        return f"OCR of image `{filename}`"
    if suffix in PDF_SUFFIXES:
        return f"text extracted from PDF `{filename}`"
    if suffix in DOCX_SUFFIXES:
        return f"text extracted from Word document `{filename}`"
    return f"imported file `{filename}`"
